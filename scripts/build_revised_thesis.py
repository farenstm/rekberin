from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SOURCE = Path(r"C:\Users\rens\Downloads\SKRIPSI_FORMATTED (23).docx")
OUTPUT = Path(r"C:\eescrow\SKRIPSI_REVISI_SUBSTANSI.docx")

CONTRACT_ADDRESS = "0x1eCB0A2Ad4495a1B050B519b6ACe92B1e068Bf92"
PRODUCTION_URL = "https://rekberin-jade.vercel.app"

REPLACEMENTS = {
    "agar bahwa": "memastikan bahwa",
    "supaya keamanan dana pengguna": "menjaga keamanan dana pengguna",
    "bisa aplikasi frontend": "memungkinkan aplikasi frontend",
    "membuat Pembeli bisa Pembeli": "memungkinkan Pembeli",
    "Fungsi ini dengan validasi state": "Fungsi ini dilengkapi validasi state",
    "fungsi ini dengan validasi state": "fungsi ini dilengkapi validasi state",
    "integrasi system": "integrasi sistem",
    "integrasi antar-komponen": "integrasi antarkomponen",
    "user experience": "pengalaman pengguna",
    "di-hold": "ditahan",
    "seller approve atau reject": "penjual menyetujui atau menolak",
    "status escrow ditampilkan secara real-time melalui pembacaan event EscrowStateChanged":
        "status escrow diperbarui setelah transaksi terkonfirmasi melalui sinkronisasi ulang data smart contract",
    "metadata diunggah ke IPFS melalui gateway Pinata":
        "frontend mengirim metadata ke API Route Vercel, kemudian API Route mengunggah metadata ke IPFS melalui layanan Pinata",
    "Penjual mengunggah metadata aset digital ke IPFS":
        "Penjual mengirim metadata melalui frontend dan API Route Vercel untuk diunggah ke IPFS",
    "CID dihasilkan dan terisi di form createListing":
        "CID valid dihasilkan oleh Pinata dan diteruskan ke fungsi createListing",
}


TEST_RESULTS = [
    ("TC-01", "Membuat listing dengan harga dan CID valid", "ListingCreated dipancarkan dan listing aktif", "Lulus"),
    ("TC-02", "Membuat listing dengan harga nol", "Transaksi ditolak", "Lulus"),
    ("TC-03", "Membuat listing dengan CID kosong", "Transaksi ditolak", "Lulus"),
    ("TC-04", "Penjual memperbarui listing aktif", "ListingUpdated dipancarkan", "Lulus"),
    ("TC-05", "Penjual membatalkan listing aktif", "ListingCancelled dipancarkan", "Lulus"),
    ("TC-06", "Pembeli membuat escrow", "EscrowCreated, state HELD, dana terkunci", "Lulus"),
    ("TC-07", "Pembayaran tidak sesuai harga", "Transaksi ditolak", "Lulus"),
    ("TC-08", "Penjual membeli listing sendiri", "Transaksi ditolak", "Lulus"),
    ("TC-09", "Listing dibeli untuk kedua kalinya", "Transaksi kedua ditolak", "Lulus"),
    ("TC-10", "Pembeli mengonfirmasi penerimaan", "HELD ke RELEASED dan dana ke penjual", "Lulus"),
    ("TC-11", "Pihak lain mengonfirmasi penerimaan", "Transaksi ditolak", "Lulus"),
    ("TC-12", "Pembeli mengajukan refund", "HELD ke REFUND_REQUESTED", "Lulus"),
    ("TC-13", "Pihak lain mengajukan refund", "Transaksi ditolak", "Lulus"),
    ("TC-14", "Penjual menyetujui refund", "REFUND_REQUESTED ke REFUNDED dan dana kembali", "Lulus"),
    ("TC-15", "Penjual menolak refund", "REFUND_REQUESTED kembali ke HELD", "Lulus"),
    ("TC-16", "Pihak lain memproses refund", "Transaksi ditolak", "Lulus"),
    ("TC-17", "Fungsi dipanggil pada state tidak valid", "Transaksi ditolak", "Lulus"),
    ("TC-18", "Aksi lanjutan setelah REFUNDED", "State terminal tidak dapat diubah", "Lulus"),
]


def iter_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def replace_text(paragraph):
    original = paragraph.text
    updated = original
    for old, new in REPLACEMENTS.items():
        updated = updated.replace(old, new)
    if updated == original:
        return False
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)
    return True


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragraph


def add_body(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(28)
    return paragraph


def fill_table(table, headers, rows):
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    table.style = "Table Grid"


def move_new_elements_before(document, marker, original_length):
    new_elements = list(document._body._body)[original_length:]
    for element in new_elements:
        marker._p.addprevious(element)


def main():
    copy2(SOURCE, OUTPUT)
    document = Document(OUTPUT)
    changed = sum(int(replace_text(paragraph)) for paragraph in iter_paragraphs(document))

    references = [p for p in document.paragraphs if p.text.strip() == "DAFTAR PUSTAKA"]
    if not references:
        raise RuntimeError("Heading DAFTAR PUSTAKA tidak ditemukan")
    marker = references[-1]
    original_length = len(document._body._body)

    document.add_page_break()
    add_heading(document, "BAB IV", 1)
    add_heading(document, "IMPLEMENTASI DAN PENGUJIAN", 1)

    add_heading(document, "4.1 Implementasi Sistem", 2)
    add_body(
        document,
        "Implementasi menghasilkan prototype marketplace akun game yang mengintegrasikan "
        "frontend React dan Next.js, MetaMask, smart contract EscrowChain, Polygon PoS Amoy "
        "Testnet, serta IPFS melalui layanan Pinata. Polygon PoS Amoy digunakan sebagai "
        "jaringan EVM-compatible untuk lingkungan pengujian. Prototype dapat diakses melalui "
        f"{PRODUCTION_URL}.",
    )

    add_heading(document, "4.1.1 Implementasi Smart Contract", 3)
    add_body(
        document,
        "Smart contract EscrowChain diimplementasikan menggunakan Solidity 0.8.28 dan "
        "OpenZeppelin ReentrancyGuard. Kontrak yang diuji pada Polygon Amoy berada pada alamat "
        f"{CONTRACT_ADDRESS}. Pemeriksaan RPC pada saat penyusunan revisi menunjukkan bytecode "
        "kontrak tersedia pada chain ID 80002. Kontrak menyimpan data Listing dan Escrow, "
        "sedangkan metadata deskriptif listing disimpan di IPFS.",
    )

    implementation_table = document.add_table(rows=1, cols=3)
    fill_table(
        implementation_table,
        ["Komponen", "Implementasi", "Keterangan"],
        [
            ("Frontend", "Next.js 16 dan React 19", "Antarmuka prototype marketplace"),
            ("Wallet", "MetaMask", "Otentikasi dan penandatanganan transaksi"),
            ("Blockchain", "Polygon PoS Amoy, chain ID 80002", "Jaringan EVM-compatible untuk pengujian"),
            ("Smart contract", CONTRACT_ADDRESS, "EscrowChain Solidity 0.8.28"),
            ("Penyimpanan metadata", "IPFS melalui Pinata", "CID disimpan pada smart contract"),
            ("Hosting", "Vercel", PRODUCTION_URL),
        ],
    )

    add_heading(document, "4.1.2 Implementasi Integrasi IPFS", 3)
    add_body(
        document,
        "Pengunggahan gambar dan metadata menggunakan alur frontend ke API Route Vercel, "
        "kemudian API Route mengirim data ke Pinata/IPFS. Kredensial Pinata disimpan sebagai "
        "environment variable pada sisi server dan tidak dikirim ke browser. Setelah unggahan "
        "berhasil, Pinata menghasilkan CID metadata. CID tersebut diteruskan oleh frontend ke "
        "fungsi createListing atau updateListing pada smart contract.",
    )

    add_heading(document, "4.1.3 Sinkronisasi Data Antarmuka", 3)
    add_body(
        document,
        "Antarmuka memperbarui data dengan menjalankan sinkronisasi ulang listing dan escrow "
        "setelah transaksi terkonfirmasi. Implementasi tidak menggunakan subscription event "
        "blockchain permanen. Event smart contract tetap digunakan sebagai bukti transaksi "
        "pada receipt dan pengujian, sedangkan tampilan memperoleh state terbaru melalui "
        "pemanggilan fungsi baca pada smart contract.",
    )

    add_heading(document, "4.2 Pengujian Smart Contract", 2)
    add_body(
        document,
        "Pengujian otomatis dilakukan menggunakan Hardhat, Chai, dan hardhat-chai-matchers "
        "pada jaringan lokal Hardhat. Perintah npm run test:contract menjalankan 18 skenario. "
        "Seluruh skenario berhasil dengan waktu eksekusi sekitar 624 milidetik pada lingkungan "
        "pengembangan. Pengujian mencakup jalur positif, validasi input, otorisasi pemanggil, "
        "transfer dana, dan transisi Finite State Machine.",
    )

    test_table = document.add_table(rows=1, cols=4)
    fill_table(test_table, ["ID", "Skenario", "Hasil Aktual", "Status"], TEST_RESULTS)

    add_heading(document, "4.3 Pengujian Integrasi", 2)
    integration_table = document.add_table(rows=1, cols=4)
    fill_table(
        integration_table,
        ["Komponen", "Metode", "Hasil Aktual", "Status"],
        [
            ("Deployment frontend", "HTTP GET production", "Website merespons HTTP 200", "Lulus"),
            ("API unggah gambar", "POST tanpa file", "API aktif dan menolak input dengan HTTP 400", "Lulus"),
            ("API metadata", "POST metadata tidak valid", "API aktif dan menolak input dengan HTTP 400", "Lulus"),
            ("Smart contract", "eth_getCode melalui RPC", "Bytecode kontrak tersedia", "Lulus"),
            ("Data listing", "nextListingId", "Tiga listing telah tercatat pada kontrak", "Lulus"),
            ("Data escrow", "nextEscrowId", "Dua escrow telah tercatat pada kontrak", "Lulus"),
        ],
    )
    add_body(
        document,
        "Pengujian transaksi testnet yang memerlukan tanda tangan dua akun MetaMask tetap harus "
        "didokumentasikan oleh peneliti dalam bentuk transaction hash dan tangkapan layar. "
        "Bukti tersebut tidak dibuat secara otomatis dalam revisi ini karena kunci privat dan "
        "persetujuan wallet berada pada pengguna.",
    )

    add_heading(document, "4.4 Analisis Hasil Pengujian", 2)
    add_body(
        document,
        "Hasil pengujian menunjukkan bahwa smart contract menerapkan transisi FSM sesuai "
        "rancangan. State HELD dapat berubah menjadi RELEASED melalui confirmReceipt atau "
        "menjadi REFUND_REQUESTED melalui requestRefund. State REFUND_REQUESTED dapat berubah "
        "menjadi REFUNDED melalui approveRefund atau kembali menjadi HELD melalui rejectRefund. "
        "Validasi pemanggil mencegah pihak yang tidak berwenang menjalankan fungsi, sedangkan "
        "validasi state menolak transisi di luar alur yang telah ditentukan.",
    )
    add_body(
        document,
        "Pengujian saldo kontrak membuktikan dana ditahan ketika escrow dibuat, diteruskan ke "
        "penjual ketika penerimaan dikonfirmasi, dan dikembalikan kepada pembeli ketika refund "
        "disetujui. Dengan demikian, pengelolaan dana on-chain dapat berjalan terpisah dari "
        "proses penyerahan akun game yang dilakukan secara off-chain.",
    )

    add_heading(document, "4.5 Keterbatasan Implementasi", 2)
    add_body(
        document,
        "Prototype tidak menyediakan arbitrator, timeout otomatis, verifikasi penyerahan akun, "
        "atau chat internal. Riwayat event blockchain juga belum direkonstruksi secara lengkap "
        "pada frontend. ERC-20 dan EIP-712 dibahas sebagai referensi teknologi, tetapi tidak "
        "diterapkan karena transaksi prototype menggunakan native MATIC dan transaksi standar "
        "MetaMask. Listing L-003 merupakan data lama dengan CID mock dan tidak digunakan sebagai "
        "bukti keberhasilan integrasi IPFS pada implementasi terbaru.",
    )

    document.add_page_break()
    add_heading(document, "BAB V", 1)
    add_heading(document, "PENUTUP", 1)
    add_heading(document, "5.1 Kesimpulan", 2)
    conclusions = [
        "Arsitektur prototype berhasil mengintegrasikan aplikasi web, MetaMask, smart contract EscrowChain, Polygon PoS Amoy Testnet, dan IPFS melalui Pinata.",
        "Smart contract berhasil mengelola pembuatan listing, penahanan dana, pelepasan dana kepada penjual, serta pengembalian dana kepada pembeli tanpa perantara manusia yang menguasai dana.",
        "Finite State Machine empat state berhasil membatasi alur transaksi menjadi HELD, RELEASED, REFUND_REQUESTED, dan REFUNDED dengan transisi yang deterministik.",
        "Delapan belas pengujian otomatis berhasil memvalidasi fungsi utama, validasi input, otorisasi aktor, transfer dana, serta penolakan terhadap transisi state yang tidak valid.",
        "Metadata listing dapat disimpan di IPFS, sedangkan smart contract hanya menyimpan CID dan data transaksi yang diperlukan, sehingga data deskriptif tidak dibebankan seluruhnya ke blockchain.",
    ]
    for conclusion in conclusions:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(conclusion)

    add_heading(document, "5.2 Saran", 2)
    suggestions = [
        "Menambahkan rekonstruksi event blockchain menggunakan filter log agar frontend dapat menampilkan audit trail transaksi yang lengkap.",
        "Menambahkan subscription event atau mekanisme polling terjadwal untuk pembaruan UI yang lebih responsif.",
        "Melakukan audit keamanan formal dan pengujian fuzzing sebelum kontrak digunakan pada jaringan produksi.",
        "Mengembangkan mekanisme penyelesaian sengketa, timeout, atau arbitrator sebagai penelitian lanjutan tanpa mengubah ruang lingkup prototype saat ini.",
        "Mengevaluasi penggunaan token ERC-20 dan EIP-712 apabila penelitian selanjutnya memerlukan stablecoin atau penandatanganan data terstruktur.",
    ]
    for suggestion in suggestions:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(suggestion)

    move_new_elements_before(document, marker, original_length)
    document.save(OUTPUT)
    print(f"output={OUTPUT}")
    print(f"redaction_changes={changed}")
    print(f"tests_documented={len(TEST_RESULTS)}")


if __name__ == "__main__":
    main()
