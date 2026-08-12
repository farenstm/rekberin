from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SOURCE = Path(r"C:\Users\rens\Downloads\SKRIPSI_FORMATTED (23).docx")
OUTPUT = Path(r"C:\eescrow\SKRIPSI_HASIL_KOREKSI.docx")


REPLACEMENTS = {
    "agar bahwa": "memastikan bahwa",
    "supaya keamanan dana pengguna": "menjaga keamanan dana pengguna",
    "bisa aplikasi frontend": "memungkinkan aplikasi frontend",
    "membuat Pembeli bisa Pembeli": "memungkinkan Pembeli",
    "Fungsi ini dengan validasi state": "Fungsi ini dilengkapi validasi state",
    "fungsi ini dengan validasi state": "fungsi ini dilengkapi validasi state",
    "integrasi system": "integrasi sistem",
    "integrasi antar-komponen": "integrasi antarkomponen",
    "user experience": "pengalaman pengguna",
    "di-hold": "ditahan",
    "seller approve atau reject": "penjual menyetujui atau menolak",
}


def iter_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                for nested in cell.tables:
                    for row2 in nested.rows:
                        for cell2 in row2.cells:
                            yield from cell2.paragraphs
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def replace_text(paragraph, replacements):
    original = paragraph.text
    updated = original
    for old, new in replacements.items():
        updated = updated.replace(old, new)
    if updated == original:
        return False

    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)
    return True


def add_note(document: Document, title: str, body: str):
    heading = document.add_paragraph()
    heading.style = document.styles["Heading 2"]
    heading.add_run(title)
    paragraph = document.add_paragraph(body)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    copy2(SOURCE, OUTPUT)
    document = Document(OUTPUT)

    changed = 0
    for paragraph in iter_paragraphs(document):
        changed += int(replace_text(paragraph, REPLACEMENTS))

    document.add_page_break()
    title = document.add_paragraph()
    title.style = document.styles["Heading 1"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CATATAN KOREKSI DAN PENYESUAIAN IMPLEMENTASI")
    run.bold = True
    run.font.size = Pt(14)

    intro = document.add_paragraph(
        "Bagian ini merupakan catatan editorial untuk penulis dan perlu dihapus sebelum "
        "naskah final dikumpulkan. Catatan disusun berdasarkan perbandingan antara draf "
        "skripsi dengan source code, smart contract, deployment Vercel, dan konfigurasi "
        "integrasi yang digunakan pada prototype saat ini."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    notes = [
        (
            "1. Lengkapi Bab IV dan Bab V",
            "Struktur naskah menyatakan lima bab, tetapi hasil ekstraksi dokumen belum "
            "menunjukkan Bab IV Implementasi dan Pengujian serta Bab V Penutup sebagai bab "
            "yang lengkap. Bab IV harus memuat hasil aktual, transaction hash, alamat smart "
            "contract, CID IPFS, tangkapan layar, dan tabel status lulus/gagal. Bab V harus "
            "memuat kesimpulan yang menjawab tujuan penelitian dan saran pengembangan.",
        ),
        (
            "2. Perbaiki dan jalankan pengujian Hardhat",
            "Test otomatis saat ini belum dapat dijadikan bukti karena perintah hardhat test "
            "masih gagal pada konfigurasi module Hardhat/ethers. Test createEscrow juga tidak "
            "boleh mengharapkan event EscrowStateChanged apabila implementasi kontrak hanya "
            "memancarkan EscrowCreated pada pembentukan escrow. Seluruh hasil dalam Bab IV "
            "harus berasal dari test yang benar-benar dijalankan dan lulus.",
        ),
        (
            "3. Sesuaikan arsitektur unggah IPFS",
            "Implementasi terbaru menggunakan alur Frontend -> API Route Vercel -> Pinata/IPFS "
            "-> CID -> Smart Contract. Frontend tidak lagi menyimpan atau mengirim secret "
            "Pinata secara langsung. Diagram arsitektur fisik, diagram komponen, sequence "
            "diagram, dan uraian write path harus menggambarkan API server tersebut.",
        ),
        (
            "4. Koreksi klaim pembaruan real-time",
            "Frontend saat ini menjalankan sinkronisasi ulang menggunakan syncListings dan "
            "syncEscrows setelah membaca smart contract. Belum terdapat subscription event "
            "blockchain real-time yang permanen. Narasi sebaiknya menyebut UI diperbarui "
            "setelah transaksi terkonfirmasi dan data disinkronkan ulang, kecuali listener "
            "event benar-benar ditambahkan ke implementasi.",
        ),
        (
            "5. Perjelas klasifikasi Polygon",
            "Gunakan istilah 'Polygon PoS Amoy Testnet sebagai jaringan EVM-compatible untuk "
            "pengujian prototype'. Penyebutan Polygon PoS sebagai Layer 2 harus disertai "
            "definisi dan sumber yang kuat karena klasifikasinya berbeda dari canonical "
            "Ethereum Layer 2 berbasis rollup.",
        ),
        (
            "6. Batasi klaim ERC-20 dan EIP-712",
            "Prototype menggunakan native MATIC melalui msg.value dan transaksi MetaMask "
            "standar. ERC-20 dan EIP-712 belum diterapkan. Keduanya dapat dipertahankan pada "
            "landasan teori atau saran pengembangan, tetapi tidak boleh dinyatakan sebagai "
            "fitur implementasi aktual.",
        ),
        (
            "7. Dokumentasikan updateListing dan cancelListing",
            "Smart contract dan frontend menyediakan fungsi updateListing dan cancelListing. "
            "Kedua fungsi harus konsisten dicantumkan pada use case Kelola Listing, tabel "
            "fungsi kontrak, event, diagram terkait, dan pengujian positif maupun negatif.",
        ),
        (
            "8. Lengkapi pengujian negatif dan otorisasi",
            "Tambahkan pengujian harga nol, CID kosong, nominal pembayaran salah, pembelian "
            "listing sendiri, pembelian ganda, pemanggil tidak berwenang, state yang tidak "
            "valid, serta sifat terminal RELEASED dan REFUNDED. Catat expected result dan "
            "actual result untuk setiap kasus.",
        ),
        (
            "9. Batasi klaim audit trail",
            "Frontend belum merekonstruksi seluruh log event historis dari blockchain; data "
            "events pada proses sinkronisasi escrow masih kosong. Jangan menyatakan bahwa "
            "prototype menyediakan audit trail event lengkap kecuali log diambil dari RPC "
            "menggunakan filter event.",
        ),
        (
            "10. Bersihkan data demonstrasi lama",
            "Listing L-003 menyimpan CID mock dari deployment lama sehingga metadata tampil "
            "Unknown. Untuk demonstrasi dan bukti Bab IV, edit listing tersebut menggunakan "
            "wallet penjual atau buat listing baru setelah integrasi Pinata yang terbaru.",
        ),
    ]

    for heading, body in notes:
        add_note(document, heading, body)

    add_note(
        document,
        "Ringkasan status koreksi otomatis",
        f"Sebanyak {changed} paragraf atau sel tabel mengalami koreksi redaksional otomatis. "
        "Koreksi substansi yang memerlukan hasil eksperimen tidak diisi dengan data rekaan; "
        "bagian tersebut ditandai melalui catatan di atas dan harus dilengkapi dari hasil "
        "pengujian aktual.",
    )

    document.save(OUTPUT)
    print(f"output={OUTPUT}")
    print(f"changed={changed}")
    print(f"size={OUTPUT.stat().st_size}")


if __name__ == "__main__":
    main()
