from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(r"C:\eescrow\SKRIPSI_REVISI_CAPTION.docx")
OUTPUT = Path(r"C:\eescrow\SKRIPSI_DRAFT_LENGKAP_BAB_1-5.docx")

SAFE_REPLACEMENTS = {
    "supaya bahwa": "memastikan bahwa",
    "peer-review": "telaah sejawat",
    "system EscrowChain": "sistem EscrowChain",
    "integrasi system": "integrasi sistem",
    "point uji": "poin pengujian",
}


def replace_paragraph(paragraph):
    original = paragraph.text
    updated = original
    for old, new in SAFE_REPLACEMENTS.items():
        updated = updated.replace(old, new)
    if updated == original:
        return False
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)
    return True


def iter_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def add_body(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(28)
    return paragraph


def enable_field_updates(document):
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


copy2(SOURCE, OUTPUT)
document = Document(OUTPUT)
changed = sum(int(replace_paragraph(p)) for p in iter_paragraphs(document))

markers = [p for p in document.paragraphs if p.text.strip() == "4.4 Analisis Hasil Pengujian"]
if not markers:
    raise RuntimeError("Marker 4.4 Analisis Hasil Pengujian tidak ditemukan")
marker = markers[-1]
original_body_length = len(document._body._body)

heading = document.add_heading("4.3.1 Pengujian Manual MetaMask dan Polygon Amoy", level=3)
add_body(
    document,
    "Pengujian manual dilakukan menggunakan dua akun MetaMask dengan peran penjual dan pembeli. "
    "Setiap transaksi harus ditandatangani oleh akun yang berwenang pada Polygon PoS Amoy Testnet. "
    "Kolom transaction hash, nomor blok, dan dokumentasi diisi menggunakan bukti aktual setelah "
    "transaksi dikonfirmasi. Placeholder pada tabel tidak boleh dipertahankan pada naskah final.",
)

caption = document.add_paragraph("Tabel 4.4 Pengujian Manual MetaMask dan Polygon Amoy")
try:
    caption.style = document.styles["CaptionTabel"]
except KeyError:
    caption.style = document.styles["Caption"]
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

table = document.add_table(rows=1, cols=6)
headers = ["ID", "Skenario", "Aktor", "Hasil yang Diharapkan", "Bukti Aktual", "Status"]
for index, value in enumerate(headers):
    table.rows[0].cells[index].text = value

rows = [
    ("IT-01", "Membuat listing dan mengunggah metadata", "Penjual", "CID valid dan ListingCreated", "[ISI CID DAN TX HASH]", "Belum dilengkapi"),
    ("IT-02", "Membeli listing", "Pembeli", "EscrowCreated dan state HELD", "[ISI TX HASH DAN BLOK]", "Belum dilengkapi"),
    ("IT-03", "Konfirmasi penerimaan", "Pembeli", "State RELEASED dan dana masuk penjual", "[ISI TX HASH DAN SALDO]", "Belum dilengkapi"),
    ("IT-04", "Mengajukan pengembalian dana", "Pembeli", "State REFUND_REQUESTED", "[ISI TX HASH DAN BLOK]", "Belum dilengkapi"),
    ("IT-05", "Menyetujui pengembalian dana", "Penjual", "State REFUNDED dan dana kembali", "[ISI TX HASH DAN SALDO]", "Belum dilengkapi"),
    ("IT-06", "Menolak pengembalian dana", "Penjual", "State kembali ke HELD", "[ISI TX HASH DAN BLOK]", "Belum dilengkapi"),
    ("IT-07", "Membuka metadata dan gambar melalui gateway", "Penjual/Pembeli", "JSON metadata dan gambar dapat diakses", "[ISI URL CID/SCREENSHOT]", "Belum dilengkapi"),
]
for values in rows:
    cells = table.add_row().cells
    for index, value in enumerate(values):
        cells[index].text = value
table.style = "Table Grid"

add_body(
    document,
    "Selain mencatat status transaksi, pengujian perlu membandingkan saldo sebelum dan sesudah "
    "pelepasan atau pengembalian dana. Biaya gas dicatat terpisah agar perubahan saldo tidak "
    "keliru ditafsirkan sebagai nilai transfer bersih. Tautan explorer menggunakan pola "
    "https://amoy.polygonscan.com/tx/{transactionHash}.",
)

new_elements = list(document._body._body)[original_body_length:]
for element in new_elements:
    marker._p.addprevious(element)

enable_field_updates(document)
document.save(OUTPUT)

print(f"output={OUTPUT}")
print(f"language_changes={changed}")
print(f"manual_test_rows={len(rows)}")
