from copy import deepcopy
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\rens\Downloads\SKRIPSI_REVISI_CAPTION_rapikan.docx")
OUTPUT = Path(r"C:\eescrow\SKRIPSI_USE_CASE_SESUAI_GAMBAR_3.7.docx")

DESIRED = [
    ("UC-01", "Create Escrow", "Buat Transaksi Escrow"),
    ("UC-02", "Confirm Receipt", "Konfirmasi Penerimaan"),
    ("UC-03", "Request Refund", "Ajukan Refund"),
    ("UC-04", "View Marketplace", "Lihat Marketplace"),
    ("UC-05", "View Detail Listing", "Lihat Detail Listing"),
    ("UC-06", "View Escrow Status", "Lihat Status Escrow"),
    ("UC-07", "Create Listing", "Create Listing"),
    ("UC-08", "Manage Listing", "Kelola Listing"),
    ("UC-09", "Approve Refund", "Setujui Refund"),
    ("UC-10", "Reject Refund", "Tolak Refund"),
]

SOURCE_NAMES = {
    "Create Escrow": "Buat Transaksi Escrow",
    "Confirm Receipt": "Konfirmasi Penerimaan",
    "Request Refund": "Ajukan Refund",
    "View Marketplace": "Lihat Marketplace",
    "View Detail Listing": "Lihat Detail Listing",
    "View Escrow Status": "Lihat Status Escrow",
    "Create Listing": "Create Listing",
    "Manage Listing": "Kelola Listing",
    "Approve Refund": "Setujui Refund",
    "Reject Refund": "Tolak Refund",
}

INTRO = {
    "Create Escrow": "Skenario Create Escrow menjelaskan proses Buyer membeli listing aktif dengan mengunci dana sebesar harga listing pada smart contract hingga transaksi diselesaikan.",
    "Confirm Receipt": "Skenario Confirm Receipt menjelaskan proses Buyer mengonfirmasi penerimaan aset digital sehingga dana escrow diteruskan kepada Seller.",
    "Request Refund": "Skenario Request Refund menjelaskan proses Buyer mengajukan pengembalian dana ketika aset digital tidak diterima atau tidak sesuai dengan kesepakatan.",
    "View Marketplace": "Skenario View Marketplace menjelaskan proses Buyer atau Seller melihat daftar listing aktif yang tersedia pada marketplace.",
    "View Detail Listing": "Skenario View Detail Listing menjelaskan proses Buyer atau Seller melihat data on-chain dan metadata IPFS dari listing yang dipilih.",
    "View Escrow Status": "Skenario View Escrow Status menjelaskan proses Buyer atau Seller memantau state dan riwayat transaksi escrow.",
    "Create Listing": "Skenario Create Listing menjelaskan proses Seller mengunggah metadata ke IPFS dan mencatat CID serta harga listing pada smart contract.",
    "Manage Listing": "Skenario Manage Listing menjelaskan proses Seller memperbarui atau membatalkan listing aktif miliknya.",
    "Approve Refund": "Skenario Approve Refund menjelaskan proses Seller menyetujui permintaan pengembalian dana dari Buyer.",
    "Reject Refund": "Skenario Reject Refund menjelaskan proses Seller menolak permintaan pengembalian dana sehingga escrow kembali ke state HELD.",
}


def set_paragraph_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text):
    cell.text = text


def find_scenario_slots(document):
    body = document._body._body
    slots = []
    for element in body:
        if element.tag != qn("w:p"):
            continue
        paragraph = Paragraph(element, document._body)
        text = paragraph.text.strip()
        if paragraph.style.name.startswith("Caption") and text.startswith("Tabel 3.") and "Skenario Use Case UC-" in text:
            next_element = element.getnext()
            while next_element is not None and next_element.tag != qn("w:tbl"):
                next_element = next_element.getnext()
            if next_element is not None:
                slots.append((paragraph, next_element))
    return slots


def scenario_name(table_element, document):
    table = Table(table_element, document._body)
    if not table.rows:
        return ""
    first = " ".join(cell.text for cell in table.rows[0].cells)
    for english, indonesian in SOURCE_NAMES.items():
        if english in first or indonesian in first:
            return english
    return ""


def replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            text = cell.text
            for old, new in replacements.items():
                text = text.replace(old, new)
            if text != cell.text:
                set_cell_text(cell, text)


def apply_implementation_corrections(name, table):
    common = {
        "Polygon Layer 2": "Polygon PoS Amoy Testnet",
        "gateway ipfs-http-client": "API Route Vercel dan layanan Pinata",
        "buyer": "Buyer",
        "seller": "Seller",
    }
    replace_in_table(table, common)

    corrections = {
        "Create Escrow": {
            '"Cannot buy own listing"': '"Seller cannot buy own listing"',
            '"Insufficient payment"': '"Incorrect MATIC value sent"',
            "Jika terjadi serangan reentrancy selama transfer, modifier nonReentrant dari OpenZeppelin ReentrancyGuard memblokir upaya tersebut [5].": "Jika transaksi tidak memenuhi validasi listing, pemilik, atau nilai pembayaran, smart contract melakukan revert sehingga dana tidak berpindah.",
        },
        "Confirm Receipt": {
            '"Not Buyer"': '"Only Buyer can confirm receipt"',
            '"Not Buyer"': '"Only Buyer can confirm receipt"',
            '"Invalid state"': '"Escrow is not HELD"',
        },
        "Request Refund": {
            '"Not Buyer"': '"Only Buyer can request refund"',
            '"Already released"': '"Escrow is not HELD"',
            '"Already requested"': '"Escrow is not HELD"',
            "atau batas waktu pengajuan refund telah lewat": "",
        },
        "Manage Listing": {
            "manageListing(listingId, newPrice, isActive)": "updateListing(listingId, newPrice, newCid) atau cancelListing(listingId)",
            '"Not listing owner"': '"Only Seller can update listing" atau "Only Seller can cancel listing"',
            '"Listing not active"': '"Listing is not active"',
            "Smart contract memperbarui data listing dan memancarkan event ListingUpdated.": "Smart contract memperbarui data dan memancarkan ListingUpdated, atau menonaktifkan listing dan memancarkan ListingCancelled.",
        },
        "Create Listing": {
            '"Invalid price or CID"': '"Price must be greater than 0" atau "CID required"',
            "node IPFS melalui API Route Vercel dan layanan Pinata": "API Route Vercel, kemudian API mengunggahnya ke Pinata/IPFS",
        },
        "Approve Refund": {
            '"Not Seller"': '"Only Seller can approve refund"',
            '"Not in refund requested state"': '"Refund not requested"',
        },
        "Reject Refund": {
            '"Not Seller"': '"Only Seller can reject refund"',
            '"Not in refund requested state"': '"Refund not requested"',
            "confirmReceipt (UC-03)": "confirmReceipt (UC-02)",
            "requestRefund (UC-04)": "requestRefund (UC-03)",
        },
        "View Marketplace": {
            "Pembeli": "Buyer atau Seller",
            "getActiveListings()": "nextListingId() dan getListing(listingId)",
        },
        "View Detail Listing": {
            "Pembeli": "Buyer atau Seller",
        },
    }
    replace_in_table(table, corrections.get(name, {}))


copy2(SOURCE, OUTPUT)
document = Document(OUTPUT)
slots = find_scenario_slots(document)
if len(slots) != 10:
    raise RuntimeError(f"Diharapkan 10 tabel skenario, ditemukan {len(slots)}")

source_xml = {}
for _, table_element in slots:
    name = scenario_name(table_element, document)
    if name:
        source_xml[name] = deepcopy(table_element)

if set(source_xml) != set(SOURCE_NAMES):
    raise RuntimeError(f"Tabel sumber tidak lengkap: {sorted(source_xml)}")

body = document._body._body
previous_table_element = None
for index, ((caption, target_element), (uc_code, name, caption_name)) in enumerate(zip(slots, DESIRED)):
    replacement = deepcopy(source_xml[name])
    target_element.addprevious(replacement)
    target_element.getparent().remove(target_element)
    slots[index] = (caption, replacement)

    set_paragraph_text(caption, f"Tabel 3.{12 + index} Skenario Use Case {uc_code} {caption_name}")
    table = Table(replacement, document._body)
    table.rows[0].cells[1].text = f"{uc_code} {name}"
    apply_implementation_corrections(name, table)

    caption_element = caption._p
    start = previous_table_element.getnext() if previous_table_element is not None else None
    if start is None:
        heading_candidates = [p for p in document.paragraphs if p.text.strip() == "3.3.4 Skenario Use Case"]
        start = heading_candidates[-1]._p.getnext()

    between = []
    current = start
    while current is not None and current is not caption_element:
        if current.tag == qn("w:p"):
            p = Paragraph(current, document._body)
            if not p.style.name.startswith("Caption"):
                between.append(p)
        current = current.getnext()
    for p in between:
        set_paragraph_text(p, "")
    if between:
        set_paragraph_text(between[-1], INTRO[name])
    previous_table_element = replacement

# Correct actor discussion: only Buyer and Seller are use-case actors.
for paragraph in document.paragraphs:
    text = paragraph.text.strip()
    if text.startswith("Pada Gambar 3.1, terlihat bahwa Penjual dan Pembeli merupakan aktor utama"):
        set_paragraph_text(
            paragraph,
            "Pada Gambar 3.7, Buyer dan Seller merupakan dua aktor yang berinteraksi langsung dengan sistem EscrowChain. MetaMask, API Vercel, Pinata/IPFS, smart contract, dan Polygon Amoy merupakan komponen pendukung yang dijelaskan pada diagram arsitektur, bukan aktor bisnis pada use case diagram.",
        )
    elif text.startswith("Tabel 3.10 menunjukkan sistem EscrowChain mengenal enam aktor"):
        set_paragraph_text(
            paragraph,
            "Tabel 3.10 menunjukkan bahwa sistem EscrowChain memiliki dua aktor utama, yaitu Seller sebagai pihak yang menawarkan aset digital dan Buyer sebagai pihak yang membeli aset melalui mekanisme escrow.",
        )

# Align actor names in the actor table with the diagram.
for table in document.tables:
    header = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
    if header == "Aktor | Tipe | Deskripsi" and len(table.rows) == 3:
        table.rows[1].cells[0].text = "Seller"
        table.rows[2].cells[0].text = "Buyer"

document.save(OUTPUT)
print(f"output={OUTPUT}")
print("actors=2")
print("use_cases=10")
print("scenario_tables=10")
